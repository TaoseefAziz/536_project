import pdfplumber
import docx
import os
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def extract_text_from_document(file):
    """
    Extracts text from a PDF or DOCX file-like object, including paragraphs, tables, and hyperlinks.

    Returns:
        The extracted text, or an error message if extraction fails or the file type is unsupported.
    """
    file_extension = os.path.splitext(file.name)[1].lower()
    text = ""

    if file_extension == ".pdf":
        try:
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

                    # Extract hyperlinks from annotations (if any)
                    if page.annots:
                        for annot in page.annots:
                            uri = annot.get("uri")
                            if uri:
                                text += f"\n[LINK] {uri}"
        except Exception as e:
            text = f"Error extracting PDF text: {str(e)}"

    elif file_extension == ".docx":
        try:
            doc = docx.Document(file)

            # Extract paragraph text
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text.strip() + "\n"

            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text += cell_text + "\n"

            # Extract hyperlinks
            for rel in doc.part.rels.values():
                if rel.reltype == RT.HYPERLINK:
                    text += f"\n[LINK] {rel.target_ref}"

        except Exception as e:
            text = f"Error extracting DOCX text: {str(e)}"

    else:
        text = "Unsupported file format. Please upload a PDF or DOCX file."

    return text
